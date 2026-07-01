"""Word (.docx) → Markdown converter using python-docx.

Walks the document element body to preserve paragraph/table/image interleaving.
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

from converters.base import BaseConverter, ConversionResult


class DOCXConverter(BaseConverter):
    """Convert Word documents to Markdown with image extraction."""

    def convert(self) -> ConversionResult:
        doc = Document(str(self.file_path))

        md_parts: list[str] = []
        images: list[tuple[str, bytes]] = []

        # Document title from first heading or filename
        title = doc.core_properties.title or ""
        if not title and doc.paragraphs:
            first = doc.paragraphs[0]
            if first.style.name.startswith("Heading"):
                title = first.text
        if title:
            md_parts.append(f"# {title}\n")

        # Walk element body for correct interleaving
        body = doc.element.body
        for child in body:
            tag = etree.QName(child).localname

            if tag == "p":
                para = self._find_paragraph(doc, child)
                if para is not None:
                    # Check for inline images
                    img_mds, pimgs = self._extract_inline_images(para, doc)
                    images.extend(pimgs)
                    if img_mds:
                        for im in img_mds:
                            md_parts.append(im)
                            md_parts.append("")

                    text = self._para_to_md(para)
                    if text:
                        md_parts.append(text)

            elif tag == "tbl":
                table = self._find_table(doc, child)
                if table is not None:
                    tbl_md = self._table_to_md(table)
                    if tbl_md:
                        md_parts.append(tbl_md)
                        md_parts.append("")

        # Collect remaining images not caught inline
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    # Check if already added
                    if img_bytes not in [b for _, b in images]:
                        filename, png_bytes = self._next_image(img_bytes, "png")
                        images.append((filename, png_bytes))
                        # Place after last paragraph
                        md_parts.append(self._image_md(filename))
                        md_parts.append("")
                except Exception:
                    continue

        metadata = {
            "format": "docx",
            "title": title,
            "author": doc.core_properties.author or "",
        }

        return ConversionResult(
            markdown=self._clean_text("\n".join(md_parts)),
            images=images,
            metadata=metadata,
        )

    # ── Paragraph → Markdown ─────────────────────────────────

    def _para_to_md(self, para) -> str:
        """Convert a single paragraph to markdown."""
        style_name = para.style.name if para.style else "Normal"

        # Skip empty paragraphs
        full_text = para.text.strip()
        if not full_text:
            return ""

        # Detect heading level
        heading_level = self._heading_level(style_name)
        if heading_level > 0:
            prefix = "#" * heading_level + " "
            return prefix + full_text

        # Detect list items
        if "List Bullet" in style_name or "ListBullet" in style_name:
            return "- " + full_text
        if "List Number" in style_name or "ListNumber" in style_name:
            # Try to get list number
            num_id = para._element.find(qn("w:pPr"))
            return "1. " + full_text

        # Rich text (bold/italic)
        if any(run.bold or run.italic for run in para.runs if run.text):
            parts = []
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"
                parts.append(text)
            return "".join(parts) + "\n"

        return full_text + "\n"

    @staticmethod
    def _heading_level(style_name: str) -> int:
        """Map Word style name to markdown heading level."""
        style_lower = style_name.lower()
        for i in range(1, 10):
            if f"heading {i}" in style_lower or f"heading{i}" in style_lower:
                return min(i, 6)
        return 0

    # ── Table → Markdown ─────────────────────────────────────

    def _table_to_md(self, table) -> str:
        """Convert a Word table to pipe-table markdown."""
        lines: list[str] = []
        max_cols = max((len(row.cells) for row in table.rows), default=0)
        if max_cols == 0:
            return ""

        is_first = True
        for row in table.rows:
            cells = [self._escape_pipe(cell.text) for cell in row.cells]
            # Pad short rows
            while len(cells) < max_cols:
                cells.append("")
            line = "| " + " | ".join(cells[:max_cols]) + " |"
            lines.append(line)
            if is_first:
                lines.append("| " + " | ".join("---" for _ in range(max_cols)) + " |")
                is_first = False

        return "\n".join(lines) + "\n"

    # ── Image extraction ─────────────────────────────────────

    def _extract_inline_images(self, para, doc) -> tuple[list[str], list[tuple[str, bytes]]]:
        """Extract images embedded in a paragraph.

        Returns:
            (list_of_markdown_refs, list_of_image_tuples)
        """
        md_refs: list[str] = []
        images: list[tuple[str, bytes]] = []

        nsmap = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }

        for blip in para._element.iter(qn("a:blip")):
            embed = blip.get(qn("r:embed"))
            if embed and embed in doc.part.rels:
                rel = doc.part.rels[embed]
                if "image" in rel.reltype:
                    try:
                        img_bytes = rel.target_part.blob
                        filename, png_bytes = self._next_image(img_bytes, "png")
                        images.append((filename, png_bytes))
                        md_refs.append(self._image_md(filename))
                    except Exception:
                        continue

        return md_refs, images

    # ── XML body helpers ─────────────────────────────────────

    def _find_paragraph(self, doc, element):
        """Find the Paragraph object matching an XML element."""
        element_id = id(element)
        for para in doc.paragraphs:
            if id(para._element) == element_id:
                return para
        return None

    def _find_table(self, doc, element):
        """Find the Table object matching an XML element."""
        element_id = id(element)
        for table in doc.tables:
            if id(table._element) == element_id:
                return table
        return None
